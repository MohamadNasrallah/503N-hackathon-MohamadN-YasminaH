"""
Build executive_brief.docx — Conut AI Operations Agent
Mirrors the Stories Coffee executive brief style.
Run: python docs/build_docx.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ── Colour palette ────────────────────────────────────────
BROWN  = RGBColor(111,  78,  55)
GOLD   = RGBColor(200, 145,  58)
CREAM  = RGBColor(247, 243, 238)
BLUE   = RGBColor( 45,  90, 142)
GREEN  = RGBColor( 39, 174,  96)
RED    = RGBColor(192,  57,  43)
WHITE  = RGBColor(255, 255, 255)
GRAY   = RGBColor(120, 120, 120)
LGRAY  = RGBColor(245, 245, 245)
BLACK  = RGBColor( 24,  27,  32)


# ── Helpers ───────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_space(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)


def _add_run(para, text, bold=False, italic=False, color=None, size=10, font_name='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = font_name
    return run


def _section_heading(doc, text):
    """Gold-underlined brown section heading matching the LaTeX style."""
    para = doc.add_paragraph()
    _para_space(para, before=10, after=2)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = BROWN
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    # Gold bottom border via paragraph border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8913A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _sidebar_box(doc, title, body, bar_color: RGBColor, bg_color: RGBColor):
    """
    A bordered box with a coloured left bar (simulated via table with 1 narrow column).
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False

    bar_col  = table.columns[0]
    body_col = table.columns[1]
    bar_col.width  = Cm(0.35)
    body_col.width = Cm(15.85)

    bar_cell  = table.cell(0, 0)
    body_cell = table.cell(0, 1)

    _set_cell_bg(bar_cell, bar_color)
    _set_cell_bg(body_cell, bg_color)

    # Remove default paragraph in bar cell
    bar_cell.paragraphs[0].clear()

    body_p = body_cell.paragraphs[0]
    body_p.clear()
    _para_space(body_p, before=2, after=0)
    if title:
        r = _add_run(body_p, title + '\n', bold=True, size=9)
    for line in body:
        p = body_cell.add_paragraph(style='Normal')
        _para_space(p, before=0, after=0)
        _add_run(p, line, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def _screenshot_placeholder(doc, title, caption):
    """A light-gray box representing a screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_bg(cell, LGRAY)

    p1 = cell.paragraphs[0]
    _para_space(p1, before=18, after=4)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p1, f'[ SCREENSHOT — {title} ]', bold=True, color=GRAY, size=9)

    p2 = cell.add_paragraph()
    _para_space(p2, before=4, after=18)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, caption, italic=True, color=GRAY, size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Document setup ────────────────────────────────────────

doc = Document()

# Narrow margins
for section in doc.sections:
    section.top_margin    = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── TITLE BLOCK ───────────────────────────────────────────
title_table = doc.add_table(rows=1, cols=1)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tc = title_table.cell(0, 0)
_set_cell_bg(tc, BROWN)

p_title = tc.paragraphs[0]
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
_para_space(p_title, before=8, after=2)
_add_run(p_title, 'Conut AI Chief of Operations Agent', bold=True, color=WHITE, size=16)

p_sub = tc.add_paragraph()
_para_space(p_sub, before=2, after=2)
_add_run(p_sub, 'CEO-ready executive brief  |  4 branches  |  5 operational objectives  |  February 2026', color=CREAM, size=9.5)

p_tag = tc.add_paragraph()
_para_space(p_tag, before=2, after=6)
_add_run(p_tag, 'EECE 503N Hackathon  |  Prepared for demo day', color=CREAM, size=8.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── QUICK-STATS STRIP ─────────────────────────────────────
stats = [
    ('4',      'BRANCHES ANALYSED'),
    ('5',      'AI MODELS'),
    ('7+',     'STAFFING ALERTS'),
    ('172.4',  'TOP LOCATION SCORE'),
    ('0 API',  'COST TO DEPLOY'),
]
stat_table = doc.add_table(rows=2, cols=5)
stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
stat_table.allow_autofit = False
col_w = Cm(3.22)
for i, col in enumerate(stat_table.columns):
    col.width = col_w

alt_bg = [RGBColor(255,248,225), CREAM, RGBColor(255,248,225), CREAM, RGBColor(255,248,225)]
for i, (val, label) in enumerate(stats):
    vc = stat_table.cell(0, i)
    lc = stat_table.cell(1, i)
    _set_cell_bg(vc, alt_bg[i])
    _set_cell_bg(lc, alt_bg[i])

    vp = vc.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(vp, before=4, after=0)
    _add_run(vp, val, bold=True, color=BROWN, size=14)

    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(lp, before=0, after=4)
    _add_run(lp, label, color=GRAY, size=7)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── PLATFORM BOX ─────────────────────────────────────────
_section_heading(doc, 'Operational Platform (live and deployable)')

platform_rows = [
    ('Dashboard:',  'Streamlit multi-page BI dashboard — 6 report pages + interactive AI agent page'),
    ('Backend API:','FastAPI with 8 REST endpoints (/combo, /demand, /expansion, /staffing, /strategy, /overview, /query, /health)'),
    ('AI Agent:',   'Local rule-based Q&A (zero API cost, works offline, deploy-ready); optionally upgrades to Google Gemini 2.5 Flash via .env — no code changes'),
    ('Models:',     'Apriori · Linear Regression · Multi-signal heuristic · Headcount buffer (×1.15) · Segment share analysis'),
    ('Data:',       '9 report-style POS CSVs auto-parsed: header stripping, copyright removal, comma-number normalisation'),
]
pt = doc.add_table(rows=len(platform_rows), cols=2)
pt.allow_autofit = False
pt.columns[0].width = Cm(2.8)
pt.columns[1].width = Cm(13.36)
for r, (key, val) in enumerate(platform_rows):
    kc = pt.cell(r, 0)
    vc = pt.cell(r, 1)
    _set_cell_bg(kc, CREAM)
    _set_cell_bg(vc, CREAM)
    kp = kc.paragraphs[0]
    _para_space(kp, before=2, after=2)
    _add_run(kp, key, bold=True, color=BROWN, size=9)
    vp = vc.paragraphs[0]
    _para_space(vp, before=2, after=2)
    _add_run(vp, val, size=9)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── PROBLEM STATEMENT ─────────────────────────────────────
_section_heading(doc, 'Problem Statement')
ps = doc.add_paragraph()
_para_space(ps, before=2, after=4)
_add_run(ps, (
    'Conut — a Lebanese sweets and beverage chain with four branches (Conut, Conut-Tyre, Conut Jnah, '
    'Main Street Coffee) — had strong brand presence but no structured system for operational '
    'decision-making. Branch managers made roster, inventory, and promotion decisions without data, '
    'while leadership lacked a single view of network health. Five specific blindspots drove unnecessary '
    'cost and lost revenue: unoptimised product bundles, no demand forecast, no expansion readiness '
    'framework, chronic understaffing on specific shifts, and underperforming beverage revenue relative '
    'to industry benchmarks. We built a complete AI-powered Operations Intelligence platform to close '
    'all five gaps simultaneously, producing a CEO-ready report and an interactive dashboard that runs '
    'on live POS data with a single command.'
), size=9.5)

# ── KEY FINDINGS ──────────────────────────────────────────
_section_heading(doc, 'Key Findings')

findings = [
    (
        'Objective 1 — Product Combo Optimization (Apriori Association Rules)',
        ('Delivery basket data mined with Apriori. 10 association rules identified; '
         '4 combos with lift ≥ 2.0× qualify as prime bundle candidates. '
         'Strongest pair: lift 3.23×, 60% cross-purchase confidence. '
         'A 5% bundle discount recovers margin within 2–3 units through basket uplift.'),
        BROWN, CREAM
    ),
    (
        'Objective 2 — Demand Forecasting by Branch (Linear Regression + Confidence Bands)',
        ('Monthly sales modelled per branch with 3–12-month projection horizon. '
         'Main Street Coffee leads at 4.49M units/month projected. '
         '60% demand gap between top and bottom branch signals uneven market penetration '
         'and asymmetric resource needs. Growing branches flagged for inventory pre-positioning.'),
        BROWN, CREAM
    ),
    (
        'Objective 3 — Expansion Feasibility (3-signal Model + 10-location Composite Scoring)',
        ('Three binary signals: network growth > 1%/month, saturation > 40% revenue at one branch, '
         'customer density > 500 network-wide. Verdict: RECOMMENDED — all 3 signals met. '
         '10 Beirut-area locations scored: demand × growth factor − (competition risk × 0.3). '
         'Top candidate: Verdun (composite score 172.4). Recommended next step: site survey.'),
        BROWN, CREAM
    ),
    (
        'Objective 4 — Shift Staffing Optimization (Mean × 1.15 Safety Buffer)',
        ('Punch-in/out logs aggregated by branch × shift × date across 7 branch-shift slots. '
         '7 at-risk shifts have recorded minimum headcounts below 2 staff. '
         'Recommended daily network total: 15 staff. Single-staff shifts create service failure, '
         'burnout, and compliance risk — requiring immediate roster correction.'),
        BROWN, CREAM
    ),
    (
        'Objective 5 — Beverage Revenue Growth (Segment Share & Benchmarking)',
        ('Combined beverage share: 5.9% (coffee 3.4% + milkshake 2.6%). '
         'Industry benchmark for specialty café-bakery: 35% — a 29.1pp gap. '
         'Underperforming branches identified at item level; '
         '5 ranked strategies generated with explicit timelines and expected impact.'),
        BROWN, CREAM
    ),
    (
        'Additive Feature — AI Operations Agent (Local Rule-based + Optional Gemini)',
        ('Natural-language Q&A embedded in the dashboard — no API key or internet required. '
         'The local agent classifies questions across all 5 domains using keyword scoring, '
         'runs the relevant model, and returns structured markdown answers with real numbers. '
         'GEMINI_API_KEY in .env seamlessly upgrades to Google Gemini 2.5 Flash; '
         '/health exposes agent_mode: local|gemini so the dashboard self-configures.'),
        BLUE, RGBColor(235, 242, 252)
    ),
]

for title, body, bar, bg in findings:
    _sidebar_box(doc, title, [body], bar, bg)

# ── SCREENSHOT ROW 1 ──────────────────────────────────────
# Two side-by-side using a 2-col table
ss1 = doc.add_table(rows=1, cols=2)
ss1.allow_autofit = False
ss1.columns[0].width = Cm(8.0)
ss1.columns[1].width = Cm(8.0)

for col_i, (ttl, cap) in enumerate([
    ('Figure 1: COO Briefing — Network Health Scorecard',
     'KPI cards + demand ranking bar chart + priority action board'),
    ('Figure 2: Combo Optimization — Apriori Lift Chart',
     'Lift bars with 2.0× threshold line + confidence vs. lift scatter'),
]):
    cell = ss1.cell(0, col_i)
    _set_cell_bg(cell, LGRAY)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p1, before=22, after=4)
    _add_run(p1, f'[ SCREENSHOT — {ttl} ]', bold=True, color=GRAY, size=8)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p2, before=4, after=22)
    _add_run(p2, cap, italic=True, color=GRAY, size=7.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── PAGE BREAK ────────────────────────────────────────────
doc.add_page_break()

# ── RECOMMENDATIONS TABLE ─────────────────────────────────
_section_heading(doc, 'Recommended Actions')

rec_rows = [
    ('1', 'This week',     'Re-roster all 7 at-risk shifts to minimum 2 staff before next operating day; use model recommended headcounts as the new baseline.',           'Service risk eliminated; compliance exposure removed', RED),
    ('2', 'This week',     'Activate top bundle (lift 3.23×) in delivery app and at POS at 5–8% combined discount.',                                                      '+15% avg delivery basket size', RED),
    ('3', 'This month',    'Launch "Coffee Drive" at underperforming branches: barista training, menu prominence, off-peak happy hour 14:00–17:00 Mon–Thu.',               '+3–5pp coffee revenue share', RGBColor(200,145,58)),
    ('4', 'This month',    'Commission Verdun site survey and lease feasibility study. Run in parallel with the beverage campaign.',                                        'Opens path to 5th branch; first-mover advantage', RGBColor(200,145,58)),
    ('5', 'This quarter',  'Milkshake seasonal SKU programme at low-share branches. Limited-edition flavours for trial without permanent menu commitment.',                 '+10–20% milkshake volume', GREEN),
    ('6', 'This quarter',  'Monthly demand review cadence: re-run forecast on new POS export and adjust inventory buffers for growing vs. declining branches.',             'Fewer stockouts; avoids over-stocking', GREEN),
]

headers = ['#', 'Timeline', 'Action', 'Expected Impact']
col_widths = [Cm(0.7), Cm(2.3), Cm(9.5), Cm(3.66)]
rec_table = doc.add_table(rows=1 + len(rec_rows), cols=4)
rec_table.allow_autofit = False
for i, w in enumerate(col_widths):
    for row in rec_table.rows:
        row.cells[i].width = w

# Header row
for i, h in enumerate(headers):
    hc = rec_table.cell(0, i)
    _set_cell_bg(hc, BROWN)
    hp = hc.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(hp, before=3, after=3)
    _add_run(hp, h, bold=True, color=WHITE, size=9)

row_bgs = [
    RGBColor(255,235,235), RGBColor(255,235,235),
    RGBColor(255,248,220), RGBColor(255,248,220),
    RGBColor(235,252,240), RGBColor(235,252,240),
]
for ri, (num, timeline, action, impact, _) in enumerate(rec_rows):
    bg = row_bgs[ri]
    for ci, text in enumerate([num, timeline, action, impact]):
        cell = rec_table.cell(ri+1, ci)
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        _para_space(p, before=3, after=3)
        if ci == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, text, bold=True, size=9)
        elif ci == 1:
            _add_run(p, text, bold=True, color=BROWN, size=9)
        else:
            _add_run(p, text, size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── EXPECTED IMPACT ───────────────────────────────────────
_section_heading(doc, 'Expected Impact')

_sidebar_box(doc,
    'If Conut executes the top actions, the principal gains are:',
    [
        'Bundle activation (lift 3.23×, 5% discount)    →  +15% avg delivery basket size',
        'Staffing correction (re-roster 7 shifts)        →  Service and compliance risk eliminated',
        'Coffee Drive (training + menu change)           →  +3–5pp beverage revenue share',
        '5th branch at Verdun (site survey → opening)   →  New revenue stream',
        'Milkshake SKU rollout at low-share branches     →  +10–20% milkshake volume',
        '',
        'Closing the 29.1pp beverage gap to the 35% industry benchmark is the single largest near-term',
        'revenue lever — requiring zero capital. Even at 50% capture rate, the beverage mix shift',
        'meaningfully improves margins given beverages command higher margins than baked goods.',
    ],
    GREEN, RGBColor(235, 252, 240)
)

# ── RISK / LIMITATIONS ────────────────────────────────────
_sidebar_box(doc,
    'Model limitations to disclose:',
    [
        '• Demand forecast: linear regression on limited time window — no seasonality without Prophet.',
        '• Expansion: calibrated heuristics, not geodemographic data.',
        '• Staffing: assumes stationarity of shift patterns — re-calibrate after roster changes.',
        '• Combo: requires minimum basket density; sparse data may fall below Apriori support threshold.',
        '• Local agent: keyword-based classification — complex multi-turn reasoning requires Gemini.',
    ],
    RED, RGBColor(252, 235, 235)
)

# ── METHODOLOGY ───────────────────────────────────────────
_section_heading(doc, 'Methodology')

meth_table = doc.add_table(rows=1, cols=2)
meth_table.allow_autofit = False
meth_table.columns[0].width = Cm(8.0)
meth_table.columns[1].width = Cm(8.16)

left_cell  = meth_table.cell(0, 0)
right_cell = meth_table.cell(0, 1)

left_items = [
    ('Data Engineering', (
        'Parsed 9 report-style POS CSV files. Each contains repeated page headers, copyright footers, '
        'comma-formatted numbers, and hierarchical metadata. Custom per-file loaders strip non-data rows, '
        'reconstruct the branch/division/group hierarchy, and produce clean analytical DataFrames.'
    )),
    ('Analytical Models', None),
]
model_bullets = [
    'Combo: Apriori (mlxtend); configurable thresholds; fallback to top-item ranking.',
    'Demand: per-branch linear regression; point forecast + 90% confidence bounds; up to 12-month horizon.',
    'Expansion: 3 binary signals; composite score = demand × growth − competition × 0.3.',
    'Staffing: mean × 1.15 per (branch, shift), capped at observed max; flags min < 2.',
    'Beverage: segment share vs. benchmarks (20% coffee / 10% shake / 35% combined).',
]

right_items = [
    ('AI Agent Architecture', (
        'Two-tier design: (1) Local agent — keyword scoring across 5 topic buckets, runs relevant model, '
        'returns structured markdown instantly, no internet. (2) Gemini 2.5 Flash — injected via .env; '
        '/query endpoint auto-selects tier; /health exposes agent_mode: local|gemini.'
    )),
    ('Productisation', None),
]
prod_bullets = [
    'FastAPI — 8 endpoints, numpy-safe JSON, CORS-enabled',
    'Streamlit — 7 pages, full CSS design system, Plotly charts',
    'One-command startup: start_api.bat + start_dashboard.bat',
    'OpenClaw skill — all endpoints exposed as a Claude Code skill',
]
repro = 'pip install -r requirements.txt\nstart_api.bat\nstart_dashboard.bat'

for cell, sections, bullets, extra in [
    (left_cell,  left_items,  model_bullets, None),
    (right_cell, right_items, prod_bullets,  repro),
]:
    for heading, body in sections:
        ph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        _para_space(ph, before=3, after=1)
        _add_run(ph, heading, bold=True, color=BROWN, size=9.5)
        if body:
            pb = cell.add_paragraph()
            _para_space(pb, before=0, after=2)
            _add_run(pb, body, size=8.5)

    for b in bullets:
        pb = cell.add_paragraph()
        _para_space(pb, before=0, after=0)
        _add_run(pb, '• ' + b, size=8.5)

    if extra:
        pe = cell.add_paragraph()
        _para_space(pe, before=4, after=0)
        _add_run(pe, 'Reproducibility', bold=True, color=BROWN, size=9.5)
        for line in extra.split('\n'):
            pl = cell.add_paragraph()
            _para_space(pl, before=0, after=0)
            _add_run(pl, line, size=8, color=BLUE)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── SCREENSHOT ROW 2 ──────────────────────────────────────
ss2_data = [
    ('Figure 3: Expansion — Signal Analysis & Location Scoring',
     '3 signal verdict cards + composite score chart + top-3 location cards'),
    ('Figure 4: Ask the Agent — Local Agent Response',
     'Natural-language question + structured markdown answer with real data'),
]
ss2 = doc.add_table(rows=1, cols=2)
ss2.allow_autofit = False
ss2.columns[0].width = Cm(8.0)
ss2.columns[1].width = Cm(8.0)
for ci, (ttl, cap) in enumerate(ss2_data):
    cell = ss2.cell(0, ci)
    _set_cell_bg(cell, LGRAY)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p1, before=22, after=4)
    _add_run(p1, f'[ SCREENSHOT — {ttl} ]', bold=True, color=GRAY, size=8)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p2, before=4, after=22)
    _add_run(p2, cap, italic=True, color=GRAY, size=7.5)

doc.add_paragraph().paragraph_format.space_after = Pt(3)

ss3_data = [
    ('Figure 5: Shift Staffing — Recommended vs. Historical Headcount',
     'Grouped bar chart + min/max range plot + at-risk alerts'),
    ('Figure 6: Beverage Growth — Segment Share & Strategy Playbook',
     'Donut chart + stacked per-branch bar + ranked strategy expanders'),
]
ss3 = doc.add_table(rows=1, cols=2)
ss3.allow_autofit = False
ss3.columns[0].width = Cm(8.0)
ss3.columns[1].width = Cm(8.0)
for ci, (ttl, cap) in enumerate(ss3_data):
    cell = ss3.cell(0, ci)
    _set_cell_bg(cell, LGRAY)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p1, before=22, after=4)
    _add_run(p1, f'[ SCREENSHOT — {ttl} ]', bold=True, color=GRAY, size=8)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p2, before=4, after=22)
    _add_run(p2, cap, italic=True, color=GRAY, size=7.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── FOOTER NOTE ───────────────────────────────────────────
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
_para_space(fp, before=4, after=0)
_add_run(fp,
    'Prepared for EECE 503N Hackathon  |  February 2026  |  '
    'All analytics run on Conut POS exports  |  No external services required',
    italic=True, color=GRAY, size=8
)

# ── SAVE ─────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), 'executive_brief.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
